"""
Text Extraction Service

This module handles extracting text content from various file formats.
Supports Google Docs, PDFs, Word documents, and Excel files.

Usage:
    extractor = TextExtractor()
    text = extractor.extract_text(file_content, mime_type)
"""

import io
import logging
from typing import Optional

import PyPDF2
from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Service for extracting text from various file formats.

    This class provides methods to convert different file types (PDF, DOCX, XLSX, etc.)
    into plain text that can be processed for embeddings and analysis.
    """

    def extract_text(self, file_content: bytes, mime_type: str, filename: str = "") -> str:
        """
        Extract text from a file based on its MIME type.

        Args:
            file_content: Raw bytes of the file
            mime_type: MIME type of the file (e.g., 'application/pdf')
            filename: Optional filename for logging purposes

        Returns:
            Extracted text as a string

        Raises:
            ValueError: If the MIME type is not supported
        """
        logger.info(f"Extracting text from {filename or 'file'} (type: {mime_type})")

        try:
            # Google Docs (exported as plain text)
            if mime_type == "text/plain":
                return self._extract_from_text(file_content)

            # PDF files
            elif mime_type == "application/pdf":
                return self._extract_from_pdf(file_content)

            # Word documents (.docx)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                return self._extract_from_docx(file_content)

            # Excel spreadsheets (.xlsx)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            ]:
                return self._extract_from_xlsx(file_content)

            # Google Docs (native format - should be exported first)
            elif mime_type == "application/vnd.google-apps.document":
                logger.warning(f"Google Doc should be exported as text/plain first: {filename}")
                return self._extract_from_text(file_content)

            # Unsupported format
            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                raise ValueError(f"Unsupported MIME type: {mime_type}")

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise

    def _extract_from_text(self, file_content: bytes) -> str:
        """
        Extract text from plain text files.

        Args:
            file_content: Raw bytes of the text file

        Returns:
            Decoded text string
        """
        try:
            # Try UTF-8 first
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            # Fallback to latin-1 if UTF-8 fails
            logger.warning("UTF-8 decode failed, trying latin-1")
            return file_content.decode("latin-1", errors="ignore")

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF files using PyPDF2.

        Args:
            file_content: Raw bytes of the PDF file

        Returns:
            Extracted text from all pages
        """
        text_parts = []
        pdf_file = io.BytesIO(file_content)

        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {str(e)}")
                    continue

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract PDF: {str(e)}")

    def _extract_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from Word documents (.docx).

        Args:
            file_content: Raw bytes of the DOCX file

        Returns:
            Extracted text from all paragraphs
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract DOCX: {str(e)}")

    def _extract_from_xlsx(self, file_content: bytes) -> str:
        """
        Extract text from Excel spreadsheets (.xlsx).

        Converts spreadsheet data into a readable text format.
        Each sheet is labeled, and cells are separated by pipes (|).

        Args:
            file_content: Raw bytes of the XLSX file

        Returns:
            Extracted text from all sheets
        """
        try:
            xlsx_file = io.BytesIO(file_content)
            workbook = load_workbook(xlsx_file, read_only=True, data_only=True)

            text_parts = []

            # Process each sheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Add sheet header
                text_parts.append(f"=== {sheet_name} ===")

                # Extract rows
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty cells and convert to strings
                    row_values = [str(cell) for cell in row if cell is not None]
                    if row_values:  # Only add non-empty rows
                        text_parts.append(" | ".join(row_values))

                text_parts.append("")  # Add blank line between sheets

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"XLSX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract XLSX: {str(e)}")

    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.

        Removes excessive whitespace, normalizes line breaks,
        and removes special characters that might cause issues.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Remove excessive whitespace
        text = " ".join(text.split())

        # Normalize line breaks (keep paragraph structure)
        text = text.replace("\r\n", "\n")
        text = text.replace("\r", "\n")

        # Remove excessive blank lines (more than 2 consecutive newlines)
        while "\n\n\n" in text:
            text = text.replace("\n\n\n", "\n\n")

        return text.strip()

    def get_word_count(self, text: str) -> int:
        """
        Count the number of words in the text.

        Args:
            text: Text to count words in

        Returns:
            Number of words
        """
        return len(text.split())
